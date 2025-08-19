import streamlit as st
import sys
import os
import re

# Configurar página ANTES que cualquier otra cosa
st.set_page_config(page_title="Generador Educativo AI", page_icon="🤖", layout="wide")

# Título principal
st.title("Generador de contenido educativo AI 🤖")
st.markdown("Genera material educativo con exportación a Word")

# Verificar imports paso a paso
with st.spinner("🔄 Verificando dependencias..."):
    # Agregar path
    try:
        sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    except Exception as e:
        st.error(f"❌ Error agregando path: {e}")

    # Verificar python-docx
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches
        from datetime import datetime
        from io import BytesIO
        DOCX_OK = True
    except ImportError as e:
        st.error(f"❌ python-docx no disponible: {e}")
        DOCX_OK = False

    # Verificar servicios Bedrock
    try:
        from core.bedrock_services import generar_programacion_curricular, generar_imagen_promocional, generar_resumen_comentarios
        SERVICES_OK = True
    except Exception as e:
        st.error(f"❌ Error importando servicios: {e}")
        SERVICES_OK = False

# Función para procesar y formatear el contenido generado
def formatear_contenido_educativo(contenido_raw, grado):
    """
    Procesa el contenido generado y lo estructura como un documento educativo profesional
    """
    # Título principal del documento
    contenido_formateado = f"""
# 📚 PROGRAMACIÓN CURRICULAR - CIENCIA Y TECNOLOGÍA

## 🎓 {grado}º DE EDUCACIÓN SECUNDARIA

---

### 📋 INFORMACIÓN GENERAL
- **Área Curricular:** Ciencia y Tecnología
- **Grado:** {grado}º de Secundaria
- **Fecha de Elaboración:** {datetime.now().strftime('%d de %B de %Y')}
- **Documento generado por:** IA Educativa

---

### 📖 CONTENIDO DE LA PROGRAMACIÓN

{contenido_raw}

---

### 📝 NOTAS METODOLÓGICAS

Esta programación curricular ha sido diseñada siguiendo los lineamientos del Currículo Nacional de la Educación Básica del Perú, adaptándose a las necesidades específicas del {grado}º grado de educación secundaria.

**Recomendaciones de implementación:**
- Considerar el contexto sociocultural de los estudiantes
- Adaptar las estrategias según los ritmos de aprendizaje
- Integrar recursos tecnológicos disponibles
- Promover el aprendizaje colaborativo y la indagación científica

---

*Documento generado automáticamente por el Sistema de IA Educativa*
"""
    return contenido_formateado

# Función mejorada para crear Word
def crear_documento_profesional(contenido, titulo, grado):
    if not DOCX_OK:
        return None
    
    doc = Document()
    
    # Configurar propiedades del documento
    doc.core_properties.title = titulo
    doc.core_properties.author = "Sistema IA Educativa"
    doc.core_properties.subject = f"Programación Curricular {grado}º Secundaria"
    
    # Título principal
    titulo_principal = doc.add_heading(f"PROGRAMACIÓN CURRICULAR", 0)
    titulo_principal.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtítulo
    subtitulo = doc.add_heading(f"CIENCIA Y TECNOLOGÍA - {grado}º SECUNDARIA", 1)
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Información del documento
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    info_run = info_para.add_run(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}\nGenerado por: IA Educativa")
    info_run.italic = True
    
    # Línea separadora
    doc.add_paragraph("=" * 80)
    
    # Procesar contenido línea por línea con formato mejorado
    lineas = contenido.split('\n')
    for line in lineas:
        line = line.strip()
        if not line:
            continue
            
        # Detectar encabezados
        if line.startswith('#'):
            level = line.count('#')
            text = line.replace('#', '').strip()
            if text:
                doc.add_heading(text, level=min(level, 3))
        
        # Detectar tablas (líneas con múltiples |)
        elif '|' in line and line.count('|') >= 2:
            # Formatear como tabla o lista
            cleaned_line = line.replace('|', ' | ').strip()
            para = doc.add_paragraph(cleaned_line)
            para.style = 'List Bullet'
        
        # Detectar listas con bullets
        elif line.startswith(('•', '-', '*', '→')):
            para = doc.add_paragraph()
            para.style = 'List Bullet'
            para.add_run(line[1:].strip())
        
        # Detectar texto en mayúsculas (posibles títulos)
        elif line.isupper() and len(line) > 5:
            doc.add_heading(line.title(), 2)
        
        # Texto normal
        else:
            if len(line) > 10:  # Solo agregar líneas con contenido significativo
                doc.add_paragraph(line)
    
    # Pie de página
    doc.add_page_break()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("GENERADO POR SISTEMA IA EDUCATIVA\nMinisterio de Educación - República del Perú")
    footer_run.italic = True
    
    # Convertir a bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# Solo mostrar tabs si todo está OK
if SERVICES_OK:
    st.success("🎉 ¡Sistema listo! Genera tu programación curricular.")
    
    # Crear tabs
    tab1, tab2, tab3 = st.tabs(["📚 Programación Curricular", "🖼️ Imágenes Educativas", "🗣️ Análisis de Comentarios"])
    
    with tab1:
        st.header("📚 Generador de Programación Curricular")
        
        # Información contextual
        with st.expander("ℹ️ Información sobre la Programación Curricular"):
            st.markdown("""
            **¿Qué incluye una programación curricular?**
            - ✅ Competencias y capacidades específicas
            - ✅ Contenidos organizados por unidades
            - ✅ Desempeños observables y medibles
            - ✅ Criterios e instrumentos de evaluación
            - ✅ Estrategias metodológicas
            
            **Basado en:** Currículo Nacional de Educación Básica - MINEDU Perú
            """)
        
        with st.form("form_prog", clear_on_submit=False):
            col1, col2 = st.columns([1, 2])
            
            with col1:
                grado = st.selectbox("🎓 Grado", [3, 4, 5], format_func=lambda x: f"{x}º Secundaria")
                
                # Contenidos predefinidos por grado
                contenidos_por_grado = {
                    3: "1. LA FÍSICA Y MAGNITUDES\n2. VECTORES EN EL PLANO\n3. CINEMÁTICA\n4. DINÁMICA LINEAL\n5. TRABAJO Y ENERGÍA\n6. QUÍMICA Y MATERIA",
                    4: "1. ONDAS Y SONIDO\n2. ÓPTICA GEOMÉTRICA\n3. ELECTRICIDAD\n4. MAGNETISMO\n5. QUÍMICA ORGÁNICA\n6. BIOQUÍMICA",
                    5: "1. FÍSICA MODERNA\n2. TERMODINÁMICA\n3. FÍSICA NUCLEAR\n4. QUÍMICA AVANZADA\n5. BIOTECNOLOGÍA\n6. INVESTIGACIÓN CIENTÍFICA"
                }
            
            with col2:
                competencia = st.text_area(
                    "🎯 Competencia Principal", 
                    "Indaga mediante métodos científicos para construir sus conocimientos.",
                    height=80,
                    help="Competencia principal del área de Ciencia y Tecnología"
                )
            
            capacidades = st.text_area(
                "⚡ Capacidades Específicas",
                "• Problematiza situaciones para hacer indagación.\n"
                "• Diseña estrategias para hacer indagación.\n" 
                "• Genera y registra datos o información.\n"
                "• Analiza datos e información.\n"
                "• Evalúa y comunica el proceso y resultados de su indagación.",
                height=120,
                help="Capacidades que desarrollará el estudiante"
            )
            
            contenidos = st.text_area(
                "📖 Contenidos Curriculares",
                contenidos_por_grado[grado],
                height=120,
                help="Contenidos organizados por unidades temáticas"
            )
            
            generar = st.form_submit_button("🎯 Generar Programación Curricular Completa", use_container_width=True)
        
        # FUERA del formulario - manejar resultados
        if generar:
            with st.spinner('🔄 Generando programación curricular profesional...'):
                try:
                    resultado_raw = generar_programacion_curricular(grado, competencia, capacidades, contenidos)
                    
                    temp = resultado_raw
                    # Formatear el contenido
                    contenido_formateado = formatear_contenido_educativo(resultado_raw, grado)
                    contenido_formateado = temp
                    st.success("✅ ¡Programación curricular generada exitosamente!")
                    
                    # Mostrar resultado formateado
                    st.markdown("---")
                    st.markdown(contenido_formateado)
                    st.markdown("---")
                    
                    # Botones de descarga
                    col1, col2, col3 = st.columns(3)
                    
                    with col1:
                        st.download_button(
                            "📄 Descargar TXT",
                            data=contenido_formateado,
                            file_name=f"programacion_curricular_{grado}to_secundaria.txt",
                            mime="text/plain",
                            key="download_txt_prog",
                            use_container_width=True
                        )
                    
                    with col2:
                        if DOCX_OK:
                            doc_bytes = crear_documento_profesional(resultado_raw, f"Programación Curricular {grado}º Secundaria", grado)
                            if doc_bytes:
                                st.download_button(
                                    "📝 Descargar WORD",
                                    data=doc_bytes,
                                    file_name=f"programacion_curricular_{grado}to_secundaria.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key="download_docx_prog",
                                    use_container_width=True
                                )
                        else:
                            st.button("📝 WORD no disponible", disabled=True, key="docx_disabled_prog", use_container_width=True)
                    
                    with col3:
                        # Botón para generar nueva programación
                        if st.button("🔄 Generar Nueva", key="nueva_prog", use_container_width=True):
                            st.rerun()
                            
                except Exception as e:
                    st.error(f"❌ Error generando programación: {str(e)}")
                    st.info("💡 Verifica la conexión con AWS Bedrock")
    
    with tab2:
        st.header("🖼️ Generador de Imágenes Educativas")
        
        with st.form("form_img"):
            col1, col2 = st.columns([2, 1])
            
            with col1:
                prompt = st.text_area(
                    "🎨 Descripción de la imagen educativa",
                    "Estudiantes de secundaria realizando un experimento de física en un laboratorio moderno, "
                    "con materiales científicos, ambiente bien iluminado, estilo fotográfico profesional",
                    height=100
                )
            
            with col2:
                st.markdown("**💡 Sugerencias:**")
                st.markdown("- Laboratorio de ciencias")
                st.markdown("- Experimento de química")
                st.markdown("- Aula de física moderna")
                st.markdown("- Estudiantes investigando")
            
            generar_img = st.form_submit_button("🎨 Generar Imagen Educativa", use_container_width=True)
        
        # FUERA del formulario
        if generar_img:
            with st.spinner('🎨 Generando imagen educativa...'):
                try:
                    imagen = generar_imagen_promocional(prompt)
                    if imagen.startswith("Error"):
                        st.error(imagen)
                    else:
                        st.subheader("🖼️ Imagen Educativa Generada")
                        st.image(imagen, caption=f"Imagen generada: {prompt[:50]}...", use_column_width=True)
                except Exception as e:
                    st.error(f"❌ Error generando imagen: {str(e)}")
    
    with tab3:
        st.header("🗣️ Análisis de Comentarios Educativos")
        
        with st.form("form_comment"):
            col1, col2 = st.columns([2, 1])
            
            with col1:
                comentarios = st.text_area(
                    "💬 Comentarios de estudiantes/docentes",
                    "Las clases de ciencia son muy interesantes y aprendo mucho.\n"
                    "Me gustaría tener más experimentos prácticos en el laboratorio.\n"
                    "A veces los conceptos de física son difíciles de entender.\n"
                    "El profesor explica muy bien los temas de química.",
                    height=120
                )
            
            with col2:
                st.markdown("**📊 El análisis incluirá:**")
                st.markdown("- Sentimientos generales")
                st.markdown("- Temas de interés")
                st.markdown("- Áreas de mejora")
                st.markdown("- Recomendaciones")
            
            analizar = st.form_submit_button("🔍 Analizar Comentarios", use_container_width=True)
        
        # FUERA del formulario
        if analizar and comentarios.strip():
            with st.spinner('🔍 Analizando comentarios educativos...'):
                try:
                    analisis_raw = generar_resumen_comentarios(comentarios)
                    
                    # Formatear análisis
                    analisis_formateado = f"""
# 📊 ANÁLISIS DE COMENTARIOS EDUCATIVOS

## 📅 {datetime.now().strftime('%d de %B de %Y')}

---

### 📝 COMENTARIOS ANALIZADOS
{comentarios}

---

### 🔍 ANÁLISIS DETALLADO
{analisis_raw}

---

### 📋 RECOMENDACIONES GENERALES
- Implementar metodologías activas de enseñanza
- Fomentar el aprendizaje experimental
- Adaptar estrategias según retroalimentación estudiantil
- Mantener comunicación constante con estudiantes

---

*Análisis generado por Sistema IA Educativa*
"""
                    
                    st.success("✅ ¡Análisis completado!")
                    st.markdown("---")
                    st.markdown(analisis_formateado)
                    st.markdown("---")
                    
                    # Botones de descarga
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.download_button(
                            "📄 Descargar Análisis TXT",
                            data=analisis_formateado,
                            file_name=f"analisis_comentarios_{datetime.now().strftime('%Y%m%d')}.txt",
                            mime="text/plain",
                            key="download_txt_analisis",
                            use_container_width=True
                        )
                    
                    with col2:
                        if DOCX_OK:
                            doc_bytes = crear_documento_profesional(analisis_raw, "Análisis de Comentarios Educativos", "Análisis")
                            if doc_bytes:
                                st.download_button(
                                    "📝 Descargar WORD", 
                                    data=doc_bytes,
                                    file_name=f"analisis_comentarios_{datetime.now().strftime('%Y%m%d')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key="download_docx_analisis",
                                    use_container_width=True
                                )
                        else:
                            st.button("📝 WORD no disponible", disabled=True, key="docx_disabled_analisis", use_container_width=True)
                            
                except Exception as e:
                    st.error(f"❌ Error en análisis: {str(e)}")
        elif analizar:
            st.warning("⚠️ Por favor ingresa algunos comentarios para analizar")

else:
    st.error("⚠️ Los servicios no están disponibles. Verifica la configuración.")
    
    with st.expander("🔧 Información de diagnóstico"):
        st.write(f"**Archivo actual:** {__file__}")
        st.write(f"**Directorio actual:** {os.getcwd()}")
        st.write(f"**Directorio del archivo:** {os.path.dirname(__file__)}")
        st.write(f"**Directorio padre:** {os.path.dirname(os.path.dirname(__file__))}")
        
        core_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'core')
        st.write(f"**Buscando core en:** {core_path}")
        st.write(f"**Core existe:** {os.path.exists(core_path)}")
        
        if os.path.exists(core_path):
            st.write(f"**Archivos en core:** {os.listdir(core_path)}")

# Footer
st.markdown("---")
st.markdown("""
<div style='text-align: center'>
<h6>🎓 Sistema de Generación de Contenido Educativo con IA</h6>
<p><em>Desarrollado para el Ministerio de Educación del Perú</em></p>
</div>
""", unsafe_allow_html=True)